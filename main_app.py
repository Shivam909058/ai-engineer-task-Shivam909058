import os
import json
import tempfile
import logging
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime
import base64
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from pinecone import Pinecone
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from langchain_text_splitters import RecursiveCharacterTextSplitter
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()
INDEX_NAME = "adgm-kb"
EMBED_MODEL = "text-embedding-3-large"
LLM_MODEL = "gpt-4-turbo-preview"
TOP_K_RETRIEVAL = 5
ADGM_CHECKLISTS = {
    "incorporation_plc": {
        "process": "Private Company Limited by Shares - Incorporation",
        "required_documents": [
            "Articles of Association",
            "Memorandum of Association",
            "Shareholder Resolution",
            "Register of Members",
            "Register of Directors",
            "Declaration of Compliance",
            "Proof of Identity (Directors/Shareholders)",
            "Proof of Address",
            "Business Plan (if Financial Services)"
        ]
    },
    "employment": {
        "process": "Employment Documentation",
        "required_documents": [
            "Employment Contract",
            "Job Description",
            "Visa Documentation (if applicable)",
            "Labor Card (if applicable)"
        ]
    }
}

class ADGMComplianceAgent:
    def __init__(self):
        self.openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        self.index = self.pc.Index(INDEX_NAME)
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100
        )
    
    def embed_text(self, text: str) -> List[float]:
        """Generate embeddings for text"""
        response = self.openai_client.embeddings.create(
            model=EMBED_MODEL,
            input=text[:8000]
        )
        return response.data[0].embedding
    
    def retrieve_context(self, query: str, filters: Dict = None) -> List[Dict]:
        """Retrieve relevant context from knowledge base"""
        try:
            query_embedding = self.embed_text(query)
            
            search_kwargs = {
                "vector": query_embedding,
                "top_k": TOP_K_RETRIEVAL,
                "include_metadata": True
            }
            
            if filters:
                search_kwargs["filter"] = filters
            
            results = self.index.query(**search_kwargs)
            
            contexts = []
            for match in results['matches']:
                contexts.append({
                    "chunk": match['metadata']['chunk'],
                    "source": match['metadata']['source'],
                    "document_type": match['metadata'].get('document_type', 'unknown'),
                    "score": match['score']
                })
            
            return contexts
        except Exception as e:
            logger.error(f"Error retrieving context: {e}")
            return []
    
    def classify_document(self, text: str) -> Tuple[str, str]:
        """Use LLM to classify document type and process"""
        prompt = f"""
        Analyze the following document text and classify it:
        
        1. Document Type (choose one):
           - Articles of Association
           - Memorandum of Association
           - Shareholder Resolution
           - Employment Contract
           - Register of Members/Directors
           - Other (specify)
        
        2. Legal Process (choose one):
           - Company Incorporation
           - Employment/HR
           - Licensing
           - Compliance/Regulatory
           - Other (specify)
        
        Document Text (first 1000 characters):
        {text[:1000]}
        
        Respond in JSON format:
        {{
            "document_type": "...",
            "legal_process": "..."
        }}
        """
        
        try:
            response = self.openai_client.chat.completions.create(
                model=LLM_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            return result.get("document_type", "Unknown"), result.get("legal_process", "Unknown")
        except Exception as e:
            logger.error(f"Error classifying document: {e}")
            return "Unknown", "Unknown"
    
    def analyze_compliance(self, text: str, doc_type: str, kb_contexts: List[Dict]) -> List[Dict]:
        """Analyze document for compliance issues - FIXED VERSION"""
        context_text = "\n\n".join([
            f"Source: {ctx['source']}\nContent: {ctx['chunk'][:300]}..."
            for ctx in kb_contexts[:3]
        ])
        
        prompt = f"""
        You are an ADGM compliance expert. Review this document section for compliance issues.
        
        Document Type: {doc_type}
        
        ADGM Requirements:
        {context_text[:2000]}
        
        Document Section to Review:
        {text[:1200]}
        
        Identify the TOP 3 most critical compliance issues. Focus on:
        1. Missing REQUIRED clauses
        2. Wrong jurisdiction (must be ADGM, not UAE Federal Courts)
        3. Non-compliant terminology
        4. Missing mandatory information
        
        IMPORTANT: You MUST respond with a valid JSON object containing an "issues" array.
        For each issue, provide the EXACT text snippet (max 50 characters) that has the problem.
        
        {{
            "issues": [
                {{
                    "section": "exact problematic text (max 50 chars)",
                    "issue": "brief description of what's wrong",
                    "severity": "High",
                    "suggestion": "specific correction needed",
                    "citation": "ADGM reference"
                }}
            ]
        }}
        
        If no issues found, return: {{"issues": []}}
        """
        
        try:
            response = self.openai_client.chat.completions.create(
                model=LLM_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                max_tokens=1000,
                response_format={"type": "json_object"}
            )
            
            content = response.choices[0].message.content
            
            try:
                parsed = json.loads(content)
                
                if isinstance(parsed, dict):
                    issues = parsed.get("issues", [])
                    if isinstance(issues, list):
                        valid_issues = []
                        for issue in issues:
                            if isinstance(issue, dict) and all(
                                key in issue for key in ["section", "issue", "severity", "suggestion", "citation"]
                            ):
                                # Ensure section is not too long
                                issue["section"] = issue["section"][:50]
                                valid_issues.append(issue)
                        return valid_issues
                
                return []
                
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error: {e}")
                return []
                
        except Exception as e:
            logger.error(f"Error in analyze_compliance: {e}")
            return []
    
    def extract_docx_text(self, file) -> Tuple[str, docx.Document]:
        """Extract text from uploaded DOCX file"""
        doc = docx.Document(file)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text), doc
    
    def insert_inline_comments(self, doc: docx.Document, issues: List[Dict], missing_docs: Dict = None, doc_metadata: Dict = None) -> str:
        """Insert inline comments with comprehensive summary page - ENHANCED VERSION"""
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
            from datetime import datetime
            
            # Create comprehensive summary page at the beginning
            self._create_summary_page(doc, issues, missing_docs, doc_metadata)
            
            # Create a mapping of section text to issues
            issue_mapping = {}
            for issue in issues:
                section_key = issue.get('section', '').strip()
                if section_key:
                    issue_mapping[section_key] = issue
            
            # Track modifications
            modifications_made = 0
            
            # Process each paragraph (starting from after our summary pages)
            for para in doc.paragraphs[10:]:  # Skip first 10 paragraphs (our summary)
                para_text = para.text.strip()
                if not para_text:
                    continue
                
                # Check if this paragraph contains any flagged text
                for section_key, issue in issue_mapping.items():
                    if section_key.lower() in para_text.lower():
                        # Add inline comment right after the problematic text
                        severity_color = {
                            "High": RGBColor(255, 0, 0),    # Red
                            "Medium": RGBColor(255, 165, 0),  # Orange
                            "Low": RGBColor(255, 255, 0)     # Yellow
                        }.get(issue.get('severity', 'Medium'), RGBColor(255, 165, 0))
                        
                        # Add the comment as a highlighted inline note
                        comment_text = f" [⚠️ {issue['severity']}: {issue['issue']} | Fix: {issue['suggestion']} | Ref: {issue['citation']}]"
                        
                        # Add the comment run
                        run = para.add_run(comment_text)
                        run.font.color.rgb = severity_color
                        run.font.size = Pt(8)
                        run.font.bold = True
                        
                        # Try to highlight the background
                        try:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        except:
                            pass  # If highlighting fails, continue without it
                        
                        modifications_made += 1
                        break  # Only add one comment per paragraph
            
            # Save the document
            output_path = tempfile.mktemp(suffix="_reviewed.docx")
            doc.save(output_path)
            
            logger.info(f"Successfully added {modifications_made} inline comments and summary page")
            return output_path
            
        except Exception as e:
            logger.error(f"Error inserting comments: {e}")
            # Save without comments if error occurs
            output_path = tempfile.mktemp(suffix="_reviewed.docx")
            doc.save(output_path)
            return output_path

    def _create_summary_page(self, doc: docx.Document, issues: List[Dict], missing_docs: Dict, doc_metadata: Dict):
        """Create a comprehensive summary page at the beginning of the document"""
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
            from datetime import datetime
            
            # Get first paragraph to insert before
            first_para = doc.paragraphs[0]
            
            # 1. Main Header
            header_para = first_para.insert_paragraph_before()
            header_para.text = "ADGM COMPLIANCE REVIEW REPORT"
            header_run = header_para.runs[0]
            header_run.font.bold = True
            header_run.font.size = Pt(16)
            header_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 2. Document Information
            doc_info_para = first_para.insert_paragraph_before()
            doc_info_text = f"""
Document: {doc_metadata.get('name', 'Unknown')}
Document Type: {doc_metadata.get('type', 'Unknown')}
Review Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
Review Status: {'⚠️ ISSUES FOUND' if issues else '✅ COMPLIANT'}
            """.strip()
            doc_info_para.text = doc_info_text
            doc_info_run = doc_info_para.runs[0]
            doc_info_run.font.size = Pt(10)
            doc_info_run.font.color.rgb = RGBColor(68, 68, 68)  # Dark gray
            
            # 3. Executive Summary
            exec_summary_para = first_para.insert_paragraph_before()
            exec_summary_para.text = "EXECUTIVE SUMMARY"
            exec_summary_run = exec_summary_para.runs[0]
            exec_summary_run.font.bold = True
            exec_summary_run.font.size = Pt(14)
            exec_summary_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # 4. Issues Summary
            if issues:
                high_issues = [i for i in issues if i.get('severity') == 'High']
                medium_issues = [i for i in issues if i.get('severity') == 'Medium']
                low_issues = [i for i in issues if i.get('severity') == 'Low']
                
                issues_summary_para = first_para.insert_paragraph_before()
                issues_summary_text = f"""
📊 COMPLIANCE ISSUES SUMMARY:
• Total Issues Found: {len(issues)}
• 🔴 High Severity: {len(high_issues)} issues
• 🟡 Medium Severity: {len(medium_issues)} issues  
• 🟢 Low Severity: {len(low_issues)} issues

⚠️  IMMEDIATE ATTENTION REQUIRED: {'Yes' if high_issues else 'No'}
                """.strip()
                issues_summary_para.text = issues_summary_text
                issues_summary_run = issues_summary_para.runs[0]
                issues_summary_run.font.size = Pt(10)
                issues_summary_run.font.color.rgb = RGBColor(153, 0, 0) if high_issues else RGBColor(0, 102, 0)
            else:
                no_issues_para = first_para.insert_paragraph_before()
                no_issues_para.text = "✅ NO COMPLIANCE ISSUES DETECTED - Document appears to be compliant with ADGM requirements."
                no_issues_run = no_issues_para.runs[0]
                no_issues_run.font.size = Pt(12)
                no_issues_run.font.bold = True
                no_issues_run.font.color.rgb = RGBColor(0, 102, 0)  # Green
            
            # 5. Missing Documents Section
            if missing_docs and missing_docs.get('missing'):
                missing_header_para = first_para.insert_paragraph_before()
                missing_header_para.text = "📋 MISSING REQUIRED DOCUMENTS"
                missing_header_run = missing_header_para.runs[0]
                missing_header_run.font.bold = True
                missing_header_run.font.size = Pt(12)
                missing_header_run.font.color.rgb = RGBColor(204, 0, 0)  # Red
                
                missing_details_para = first_para.insert_paragraph_before()
                missing_text = f"""
Process: {missing_docs.get('process', 'Unknown')}
Documents Provided: {missing_docs.get('uploaded_count', 0)} of {missing_docs.get('required_count', 0)} required

❌ MISSING DOCUMENTS:
"""
                for doc in missing_docs.get('missing', []):
                    missing_text += f"• {doc}\n"
                
                missing_text += f"""
⚠️  COMPLIANCE RISK: HIGH - Missing documents may prevent successful registration/approval.
📋 ACTION REQUIRED: Please provide the missing documents before submission to ADGM.
                """.strip()
                
                missing_details_para.text = missing_text
                missing_details_run = missing_details_para.runs[0]
                missing_details_run.font.size = Pt(9)
                missing_details_run.font.color.rgb = RGBColor(153, 0, 0)
            elif missing_docs:
                complete_docs_para = first_para.insert_paragraph_before()
                complete_docs_para.text = f"✅ DOCUMENT COMPLETENESS: All {missing_docs.get('uploaded_count', 0)} required documents provided for {missing_docs.get('process', 'this process')}."
                complete_docs_run = complete_docs_para.runs[0]
                complete_docs_run.font.size = Pt(10)
                complete_docs_run.font.color.rgb = RGBColor(0, 102, 0)
            
            # 6. Legend/Guide
            legend_para = first_para.insert_paragraph_before()
            legend_para.text = "📖 REVIEW GUIDE"
            legend_run = legend_para.runs[0]
            legend_run.font.bold = True
            legend_run.font.size = Pt(12)
            legend_run.font.color.rgb = RGBColor(0, 51, 102)
            
            guide_para = first_para.insert_paragraph_before()
            guide_text = """
This document has been reviewed for ADGM compliance. Issues are marked inline with the following format:
[⚠️ Severity: Issue Description | Fix: Suggested correction | Ref: ADGM reference]

🔴 High Severity: Critical issues that must be fixed before submission
🟡 Medium Severity: Important issues that should be addressed  
🟢 Low Severity: Minor issues or recommendations for improvement

📞 For questions about this review, contact your ADGM compliance advisor.
            """.strip()
            guide_para.text = guide_text
            guide_run = guide_para.runs[0]
            guide_run.font.size = Pt(9)
            guide_run.font.color.rgb = RGBColor(68, 68, 68)
            
            # 7. Page break before original content
            page_break_para = first_para.insert_paragraph_before()
            page_break_para.text = "─" * 80
            page_break_run = page_break_para.runs[0]
            page_break_run.font.color.rgb = RGBColor(200, 200, 200)
            
            separator_para = first_para.insert_paragraph_before()
            separator_para.text = "ORIGINAL DOCUMENT WITH INLINE COMPLIANCE COMMENTS"
            separator_run = separator_para.runs[0]
            separator_run.font.bold = True
            separator_run.font.size = Pt(12)
            separator_run.font.color.rgb = RGBColor(0, 51, 102)
            separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            first_para.insert_paragraph_before().text = ""
            
        except Exception as e:
            logger.error(f"Error creating summary page: {e}")
            # Add minimal summary if detailed version fails
            try:
                simple_summary = first_para.insert_paragraph_before()
                simple_summary.text = f"ADGM COMPLIANCE REVIEW - {len(issues)} issues found"
                simple_summary.runs[0].font.bold = True
                simple_summary.runs[0].font.size = Pt(14)
            except:
                pass

    def check_missing_documents(self, uploaded_docs: List[str], process_type: str) -> Dict:
        """Check for missing required documents"""
        checklist_key = None
        if "incorporation" in process_type.lower():
            checklist_key = "incorporation_plc"
        elif "employment" in process_type.lower():
            checklist_key = "employment"
        
        if not checklist_key or checklist_key not in ADGM_CHECKLISTS:
            return {
                "process": process_type,
                "required_count": 0,
                "uploaded_count": len(uploaded_docs),
                "missing": []
            }
        
        checklist = ADGM_CHECKLISTS[checklist_key]
        required = checklist["required_documents"]
        
        missing = []
        for req_doc in required:
            found = False
            for uploaded in uploaded_docs:
                if req_doc.lower() in uploaded.lower():
                    found = True
                    break
            if not found:
                missing.append(req_doc)
        
        return {
            "process": checklist["process"],
            "required_count": len(required),
            "uploaded_count": len(uploaded_docs),
            "missing": missing
        }
    
    def process_document(self, file, missing_docs: Dict = None) -> Dict:
        """Process a single document through the compliance pipeline - ENHANCED"""
        try:
            # Extract text and document object
            text, doc = self.extract_docx_text(file)
            
            # Classify document
            doc_type, process_type = self.classify_document(text)
            
            # Split into sections for analysis
            sections = self.splitter.split_text(text)
            
            all_issues = []
            
            # Analyze each section (limit to first 8 for performance)
            for i, section in enumerate(sections[:8]):
                if len(section.strip()) < 30:  # Skip very short sections
                    continue
                
                # Retrieve relevant ADGM context
                filters = {"document_type": doc_type.lower()} if doc_type != "Unknown" else None
                contexts = self.retrieve_context(section, filters)
                
                # Analyze compliance
                issues = self.analyze_compliance(section, doc_type, contexts)
                
                # Add metadata to issues
                for issue in issues:
                    issue["section_index"] = i
                    issue["document"] = file.name
                
                all_issues.extend(issues)
            
            # Prepare document metadata
            doc_metadata = {
                'name': file.name,
                'type': doc_type,
                'process': process_type
            }
            
            # Insert comments in document with enhanced summary
            reviewed_path = self.insert_inline_comments(doc, all_issues, missing_docs, doc_metadata)
            
            return {
                "document": file.name,
                "document_type": doc_type,
                "process_type": process_type,
                "issues": all_issues,
                "reviewed_path": reviewed_path
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file.name}: {e}")
            return {
                "document": file.name,
                "error": str(e),
                "issues": []
            }

def get_file_download_link(file_path: str, filename: str) -> str:
    """Generate a download link for a file"""
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        b64 = base64.b64encode(data).decode()
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download {filename}</a>'
    except:
        return f"Error generating download link for {filename}"

def display_document_preview(file_path: str, max_paragraphs: int = 10):
    """Display a preview of the processed document"""
    try:
        doc = docx.Document(file_path)
        preview_text = []
        
        for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
            if para.text.strip():
                preview_text.append(para.text.strip())
        
        if len(doc.paragraphs) > max_paragraphs:
            preview_text.append(f"\n... ({len(doc.paragraphs) - max_paragraphs} more paragraphs)")
        
        return "\n\n".join(preview_text)
    except Exception as e:
        return f"Error loading preview: {str(e)}"

def main():
    st.set_page_config(
        page_title="ADGM Corporate Agent",
        page_icon="🏛️",
        layout="wide"
    )
    
    st.title("🏛️ ADGM Corporate Agent")
    st.markdown("""
    Upload your legal documents for automated ADGM compliance review.
    The system will:
    - ✅ Verify document completeness against ADGM checklists
    - 🔍 Detect compliance issues and red flags
    - 💬 Add **inline comments directly in the relevant clauses**
    - 📊 Generate detailed JSON reports
    """)
    
    # Initialize session state
    if 'results' not in st.session_state:
        st.session_state.results = []
    if 'all_issues' not in st.session_state:
        st.session_state.all_issues = []
    if 'missing_docs' not in st.session_state:
        st.session_state.missing_docs = {}
    if 'agent' not in st.session_state:
        with st.spinner("Initializing ADGM Compliance Agent..."):
            st.session_state.agent = ADGMComplianceAgent()
    
    agent = st.session_state.agent
    
    # Reset button
    if st.button("🔄 Reset Analysis", type="secondary"):
        st.session_state.results = []
        st.session_state.all_issues = []
        st.session_state.missing_docs = {}
        st.success("Analysis reset! You can now upload new documents.")
        st.rerun()
    
    # File upload
    uploaded_files = st.file_uploader(
        "Upload your .docx documents",
        type=["docx"],
        accept_multiple_files=True,
        help="Upload all documents related to your ADGM process (e.g., AoA, MoA, Resolutions)"
    )
    
    if uploaded_files and st.button("🚀 Run Compliance Review", type="primary"):
        # Clear previous results
        st.session_state.results = []
        st.session_state.all_issues = []
        
        # Process documents
        doc_names = []
        process_types = set()
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # First, get missing documents info for all files
        temp_doc_names = [file.name for file in uploaded_files]
        temp_process = "Company Incorporation"  # Default, will be updated
        missing_docs_info = agent.check_missing_documents(temp_doc_names, temp_process)
        
        for i, file in enumerate(uploaded_files):
            status_text.text(f"Processing {file.name}...")
            progress_bar.progress((i + 1) / len(uploaded_files))
            
            # Process document with missing docs info
            result = agent.process_document(file, missing_docs_info)
            st.session_state.results.append(result)
            
            if "error" not in result:
                st.session_state.all_issues.extend(result["issues"])
                doc_names.append(result["document"])
                process_types.add(result["process_type"])
        
        # Update missing docs with actual process type
        main_process = list(process_types)[0] if process_types else "Unknown"
        st.session_state.missing_docs = agent.check_missing_documents(doc_names, main_process)
        
        progress_bar.empty()
        status_text.empty()
        
        st.success("✅ Analysis complete! Each document now includes a comprehensive summary page.")
    
    # Display results if available
    if st.session_state.results:
        # Summary metrics
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Documents Uploaded", len(st.session_state.results))
        
        with col2:
            st.metric("Required Documents", st.session_state.missing_docs.get("required_count", 0))
        
        with col3:
            st.metric("Total Issues Found", len(st.session_state.all_issues))
        
        # Missing documents alert
        missing = st.session_state.missing_docs.get("missing", [])
        if missing:
            st.error(f"⚠️ Missing {len(missing)} required documents:")
            for doc in missing:
                st.write(f"• {doc}")
        
        # Issues summary
        if st.session_state.all_issues:
            st.warning(f"🔍 Found {len(st.session_state.all_issues)} compliance issues")
            
            # Group issues by severity
            high_issues = [i for i in st.session_state.all_issues if i.get("severity") == "High"]
            medium_issues = [i for i in st.session_state.all_issues if i.get("severity") == "Medium"]
            low_issues = [i for i in st.session_state.all_issues if i.get("severity") == "Low"]
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("🔴 High Severity", len(high_issues))
            with col2:
                st.metric("🟡 Medium Severity", len(medium_issues))
            with col3:
                st.metric("🟢 Low Severity", len(low_issues))
        
        # Download section
        st.subheader("📥 Download Results")
        
        for result in st.session_state.results:
            if "error" not in result and "reviewed_path" in result:
                st.write(f"**{result['document']}** ({result['document_type']})")
                
                col1, col2, col3 = st.columns([2, 2, 3])
                
                with col1:
                    with open(result["reviewed_path"], "rb") as f:
                        st.download_button(
                            f"📄 Download Reviewed",
                            f.read(),
                            file_name=f"reviewed_{result['document']}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{result['document']}"
                        )
                
                with col2:
                    # Generate JSON report for this document
                    doc_report = {
                        "document": result["document"],
                        "document_type": result["document_type"],
                        "issues_count": len(result["issues"]),
                        "issues": result["issues"]
                    }
                    
                    st.download_button(
                        f"📊 JSON Report",
                        json.dumps(doc_report, indent=2),
                        file_name=f"{result['document']}_report.json",
                        mime="application/json",
                        key=f"json_{result['document']}"
                    )
                
                with col3:
                    # Preview toggle
                    if st.button(f"👁️ Preview", key=f"preview_{result['document']}"):
                        st.session_state[f"show_preview_{result['document']}"] = not st.session_state.get(f"show_preview_{result['document']}", False)
                
                # Show preview if toggled
                if st.session_state.get(f"show_preview_{result['document']}", False):
                    with st.expander(f"📄 Preview: {result['document']}", expanded=True):
                        preview_text = display_document_preview(result["reviewed_path"])
                        st.text_area(
                            "Document Preview (with inline comments)",
                            preview_text,
                            height=300,
                            key=f"preview_area_{result['document']}"
                        )
        
        # Combined JSON report - FIXED VERSION
        st.subheader("📊 Combined Report")
        if st.button("📊 Generate Combined Report", key="combined_report_btn"):
            try:
                main_process = st.session_state.missing_docs.get("process", "Unknown Process")
                
                # Safely count issues by severity
                high_issues = [i for i in st.session_state.all_issues if str(i.get("severity", "")).upper() == "HIGH"]
                medium_issues = [i for i in st.session_state.all_issues if str(i.get("severity", "")).upper() == "MEDIUM"]
                low_issues = [i for i in st.session_state.all_issues if str(i.get("severity", "")).upper() == "LOW"]
                
                combined_report = {
                    "review_date": datetime.now().isoformat(),
                    "process": main_process,
                    "documents_uploaded": len(st.session_state.results),
                    "required_documents": st.session_state.missing_docs.get("required_count", 0),
                    "missing_documents": st.session_state.missing_docs.get("missing", []),
                    "total_issues": len(st.session_state.all_issues),
                    "issues_by_severity": {
                        "high": len(high_issues),
                        "medium": len(medium_issues),
                        "low": len(low_issues)
                    },
                    "issues": [
                        {
                            "document": str(issue.get("document", "Unknown")),
                            "section": str(issue.get("section", ""))[:100],
                            "issue": str(issue.get("issue", "No description")),
                            "severity": str(issue.get("severity", "Medium")),
                            "suggestion": str(issue.get("suggestion", "")),
                            "citation": str(issue.get("citation", ""))
                        }
                        for issue in st.session_state.all_issues
                    ],
                    "documents_reviewed": [
                        {
                            "name": str(r.get("document", "Unknown")),
                            "type": str(r.get("document_type", "Unknown")),
                            "issues_count": len(r.get("issues", []))
                        }
                        for r in st.session_state.results if "error" not in r
                    ]
                }
                
                # Convert to JSON
                json_output = json.dumps(combined_report, indent=2, ensure_ascii=False)
                
                st.download_button(
                    "💾 Download Combined Report",
                    json_output,
                    file_name=f"adgm_compliance_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    key="combined_download"
                )
                
                st.success("✅ Combined report ready for download!")
                
                # Show report summary
                with st.expander("📋 Report Summary", expanded=True):
                    st.json(combined_report)
                
            except Exception as e:
                st.error(f"❌ Error generating combined report: {str(e)}")
                logger.error(f"Combined report error: {e}", exc_info=True)
        
        # Detailed issues view
        if st.checkbox("🔍 Show Detailed Issues"):
            for issue in st.session_state.all_issues:
                severity_emoji = {
                    "High": "🔴",
                    "Medium": "🟡", 
                    "Low": "🟢"
                }.get(issue.get("severity", "Medium"), "🟡")
                
                with st.expander(f"{severity_emoji} {issue.get('document', 'Unknown')} - {issue.get('issue', 'No description')[:50]}..."):
                    st.write(f"**Section:** {issue.get('section', 'N/A')}")
                    st.write(f"**Issue:** {issue.get('issue', 'N/A')}")
                    st.write(f"**Severity:** {issue.get('severity', 'N/A')}")
                    st.write(f"**Suggestion:** {issue.get('suggestion', 'N/A')}")
                    st.write(f"**Reference:** {issue.get('citation', 'N/A')}")

if __name__ == "__main__":
    main()